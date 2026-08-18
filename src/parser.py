"""
Resume Screening Agent
======================

Document ingestion and parsing layer.

Responsibilities:
    - Parse TXT, PDF, and DOCX resumes.
    - Parse the Job Description.
    - Validate input files.
    - Normalize extracted text.
    - Load multiple resumes from a directory.
    - Return structured document data for downstream NLP processing.

This module intentionally contains no scoring or ranking logic.
Those responsibilities belong to separate components.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader


# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = frozenset({".txt", ".pdf", ".docx"})

DEFAULT_ENCODING = "utf-8"


# ---------------------------------------------------------------------------
# Data Models
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class ParsedDocument:
    """
    Represents a successfully parsed document.

    Attributes:
        filename: Original document filename.
        file_path: Absolute path to the document.
        file_type: Document extension.
        text: Cleaned extracted text.
        character_count: Number of extracted characters.
    """

    filename: str
    file_path: Path
    file_type: str
    text: str
    character_count: int


# ---------------------------------------------------------------------------
# Exceptions
# ---------------------------------------------------------------------------

class ParserError(Exception):
    """Base exception for document parsing failures."""


class UnsupportedFileTypeError(ParserError):
    """Raised when a document format is not supported."""


class DocumentExtractionError(ParserError):
    """Raised when text extraction fails."""


# ---------------------------------------------------------------------------
# Text Normalization
# ---------------------------------------------------------------------------

def normalize_text(text: str) -> str:
    """
    Normalize extracted document text.

    Operations:
        - Normalize line endings.
        - Remove excessive whitespace.
        - Preserve paragraph boundaries.
        - Remove accidental repeated blank lines.

    Args:
        text: Raw extracted document text.

    Returns:
        Cleaned document text.
    """

    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Remove trailing spaces from individual lines.
    text = "\n".join(line.strip() for line in text.splitlines())

    # Collapse excessive spaces/tabs.
    text = re.sub(r"[ \t]+", " ", text)

    # Keep meaningful paragraph breaks while removing excessive blank lines.
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


# ---------------------------------------------------------------------------
# TXT Parser
# ---------------------------------------------------------------------------

def _parse_txt(file_path: Path) -> str:
    """Extract text from a TXT document."""

    try:
        return file_path.read_text(
            encoding=DEFAULT_ENCODING
        )

    except UnicodeDecodeError as exc:
        raise DocumentExtractionError(
            f"Unable to decode text file: {file_path.name}"
        ) from exc

    except OSError as exc:
        raise DocumentExtractionError(
            f"Unable to read text file: {file_path.name}"
        ) from exc


# ---------------------------------------------------------------------------
# PDF Parser
# ---------------------------------------------------------------------------

def _parse_pdf(file_path: Path) -> str:
    """Extract text from a PDF document."""

    try:
        reader = PdfReader(str(file_path))

        pages = []

        for page_number, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""

                if page_text.strip():
                    pages.append(page_text)

            except Exception as exc:
                logger.warning(
                    "Failed to extract page %s from %s: %s",
                    page_number,
                    file_path.name,
                    exc,
                )

        return "\n\n".join(pages)

    except Exception as exc:
        raise DocumentExtractionError(
            f"Unable to parse PDF: {file_path.name}"
        ) from exc


# ---------------------------------------------------------------------------
# DOCX Parser
# ---------------------------------------------------------------------------

def _parse_docx(file_path: Path) -> str:
    """Extract text from a DOCX document."""

    try:
        document = Document(str(file_path))

        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        return "\n".join(paragraphs)

    except Exception as exc:
        raise DocumentExtractionError(
            f"Unable to parse DOCX: {file_path.name}"
        ) from exc


# ---------------------------------------------------------------------------
# Document Parser
# ---------------------------------------------------------------------------

def parse_document(file_path: str | Path) -> ParsedDocument:
    """
    Parse a supported document and return structured extracted data.

    Supported formats:
        .txt
        .pdf
        .docx

    Args:
        file_path: Path to the document.

    Returns:
        ParsedDocument instance.

    Raises:
        FileNotFoundError:
            If the document does not exist.

        UnsupportedFileTypeError:
            If the file format is unsupported.

        DocumentExtractionError:
            If extraction fails.
    """

    path = Path(file_path).resolve()

    if not path.exists():
        raise FileNotFoundError(
            f"Document not found: {path}"
        )

    if not path.is_file():
        raise ParserError(
            f"Expected a file but received: {path}"
        )

    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{extension}'. "
            f"Supported formats: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    logger.info(
        "Parsing document: %s",
        path.name,
    )

    if extension == ".txt":
        raw_text = _parse_txt(path)

    elif extension == ".pdf":
        raw_text = _parse_pdf(path)

    elif extension == ".docx":
        raw_text = _parse_docx(path)

    else:
        raise UnsupportedFileTypeError(
            f"Unsupported document type: {extension}"
        )

    cleaned_text = normalize_text(raw_text)

    if not cleaned_text:
        raise DocumentExtractionError(
            f"No readable text found in: {path.name}"
        )

    return ParsedDocument(
        filename=path.name,
        file_path=path,
        file_type=extension,
        text=cleaned_text,
        character_count=len(cleaned_text),
    )


# ---------------------------------------------------------------------------
# Resume Loader
# ---------------------------------------------------------------------------

def load_resumes(resume_directory: str | Path) -> list[ParsedDocument]:
    """
    Load and parse all supported resumes from a directory.

    Invalid documents are logged and skipped instead of stopping
    the entire batch-processing operation.

    Args:
        resume_directory: Directory containing resume files.

    Returns:
        List of successfully parsed resumes.
    """

    directory = Path(resume_directory).resolve()

    if not directory.exists():
        raise FileNotFoundError(
            f"Resume directory not found: {directory}"
        )

    if not directory.is_dir():
        raise NotADirectoryError(
            f"Expected a directory: {directory}"
        )

    documents: list[ParsedDocument] = []

    for file_path in sorted(directory.iterdir()):

        if not file_path.is_file():
            continue

        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            logger.debug(
                "Skipping unsupported file: %s",
                file_path.name,
            )
            continue

        try:
            document = parse_document(file_path)
            documents.append(document)

            logger.info(
                "Successfully parsed: %s",
                file_path.name,
            )

        except ParserError as exc:
            logger.error(
                "Failed to parse %s: %s",
                file_path.name,
                exc,
            )

    logger.info(
        "Resume ingestion completed: %d documents loaded",
        len(documents),
    )

    return documents


# ---------------------------------------------------------------------------
# Job Description Loader
# ---------------------------------------------------------------------------

def load_job_description(
    file_path: str | Path,
) -> ParsedDocument:
    """
    Load and parse the Job Description.

    Args:
        file_path: Path to the Job Description.

    Returns:
        ParsedDocument containing the Job Description text.
    """

    return parse_document(file_path)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

__all__ = [
    "ParsedDocument",
    "ParserError",
    "UnsupportedFileTypeError",
    "DocumentExtractionError",
    "normalize_text",
    "parse_document",
    "load_resumes",
    "load_job_description",
]


# ---------------------------------------------------------------------------
# Local Development Test
# ---------------------------------------------------------------------------

if __name__ == "__main__":

    logging.basicConfig(
        level=logging.INFO,
        format="%(levelname)s | %(message)s",
    )

    print("\n" + "=" * 70)
    print("🤖 RESUME SCREENING AGENT")
    print("📄 DOCUMENT INGESTION TEST")
    print("=" * 70)

    job_description = load_job_description(
        "data/job_description.txt"
    )

    resumes = load_resumes(
        "data/resumes"
    )

    print("\n📌 Job Description")
    print(f"   File       : {job_description.filename}")
    print(f"   Characters : {job_description.character_count}")

    print("\n📄 Resumes")
    print(f"   Loaded     : {len(resumes)}")

    for index, resume in enumerate(resumes, start=1):
        print(
            f"   {index:02d}. "
            f"{resume.filename:<25} "
            f"{resume.character_count:>6} characters"
        )

    print("\n" + "=" * 70)
    print("✅ DOCUMENT INGESTION COMPLETED")
    print("=" * 70)